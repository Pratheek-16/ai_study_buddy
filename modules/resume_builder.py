import time
import re
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER 
import io


def polish_bullets(raw_text: str, section: str, client, model: str) -> list[str]:
    """
    Takes rough user input for a section and returns polished professional bullet points.
    """
    prompt = f"""
You are a professional resume writer. The user has provided rough notes for their resume's {section} section.
Convert this into 2-4 concise, impactful bullet points using strong action verbs and quantifiable results where possible.
Return ONLY the bullet points, one per line, starting each with a bullet character •.
No intro text, no explanations, no markdown formatting.

User's notes:
{raw_text}
"""
    for attempt in range(3):
        try:
            response = client.models.generate_content(model=model, contents=prompt)
            lines = [l.strip().lstrip("•-").strip() for l in response.text.strip().split("\n") if l.strip()]
            return [l for l in lines if l]
        except Exception:
            if attempt < 2:
                time.sleep(2)
            continue
    return [raw_text]


def generate_resume_pdf(data: dict) -> bytes:
    """
    Generates a clean, ATS-friendly PDF resume from the structured data dict.
    Returns raw PDF bytes.
    """
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18*mm,
        rightMargin=18*mm,
        topMargin=16*mm,
        bottomMargin=16*mm
    )

    # ── Color palette ──────────────────────────────────────────────────────────
    dark = colors.HexColor("#1C2C25")
    accent = colors.HexColor("#2E7D52")
    soft = colors.HexColor("#555555")
    line_color = colors.HexColor("#CCCCCC")

    # ── Styles ──────────────────────────────────────────────────────────────────
    styles = getSampleStyleSheet()

    name_style = ParagraphStyle(
        "Name", fontSize=22, leading=26, textColor=dark,
        fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=2
    )
    contact_style = ParagraphStyle(
        "Contact", fontSize=9, leading=13, textColor=soft,
        fontName="Helvetica", alignment=TA_CENTER, spaceAfter=6
    )
    section_header_style = ParagraphStyle(
        "SectionHeader", fontSize=10, leading=14, textColor=accent,
        fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=3,
        textTransform="uppercase", letterSpacing=1
    )
    body_style = ParagraphStyle(
        "Body", fontSize=9.5, leading=14, textColor=dark,
        fontName="Helvetica", spaceAfter=2
    )
    bullet_style = ParagraphStyle(
        "Bullet", fontSize=9.5, leading=14, textColor=dark,
        fontName="Helvetica", leftIndent=12, spaceAfter=1,
        bulletIndent=4
    )
    sub_style = ParagraphStyle(
        "Sub", fontSize=9, leading=13, textColor=soft,
        fontName="Helvetica-Oblique", spaceAfter=2
    )

    story = []

    def section(title):
        story.append(Paragraph(title, section_header_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=line_color, spaceAfter=4))

    def bullets(items):
        for item in items:
            story.append(Paragraph(f"• {item}", bullet_style))

    # ── Name + Contact ──────────────────────────────────────────────────────────
    p = data.get("personal", {})
    story.append(Paragraph(p.get("name", ""), name_style))

    contact_parts = [x for x in [
        p.get("email", ""),
        p.get("phone", ""),
        p.get("linkedin", ""),
        p.get("location", "")
    ] if x.strip()]
    story.append(Paragraph("  |  ".join(contact_parts), contact_style))
    story.append(HRFlowable(width="100%", thickness=1, color=accent, spaceAfter=6))

    # ── Summary ─────────────────────────────────────────────────────────────────
    if data.get("summary"):
        section("Professional Summary")
        story.append(Paragraph(data["summary"], body_style))
        story.append(Spacer(1, 4))

    # ── Education ───────────────────────────────────────────────────────────────
    if data.get("education"):
        section("Education")
        for edu in data["education"]:
            story.append(Paragraph(
                f"<b>{edu.get('degree', '')}</b> — {edu.get('institution', '')}",
                body_style
            ))
            story.append(Paragraph(
                f"{edu.get('year', '')}  {('| CGPA/GPA: ' + edu.get('gpa','')) if edu.get('gpa') else ''}",
                sub_style
            ))

    # ── Experience ──────────────────────────────────────────────────────────────
    if data.get("experience"):
        section("Experience")
        for exp in data["experience"]:
            story.append(Paragraph(
                f"<b>{exp.get('role', '')}</b> — {exp.get('company', '')}",
                body_style
            ))
            story.append(Paragraph(exp.get("duration", ""), sub_style))
            bullets(exp.get("bullets", []))
            story.append(Spacer(1, 4))

    # ── Projects ─────────────────────────────────────────────────────────────────
    if data.get("projects"):
        section("Projects")
        for proj in data["projects"]:
            story.append(Paragraph(
                f"<b>{proj.get('name', '')}</b>"
                + (f" — <i>{proj.get('tech', '')}</i>" if proj.get("tech") else ""),
                body_style
            ))
            bullets(proj.get("bullets", []))
            story.append(Spacer(1, 4))

    # ── Skills ───────────────────────────────────────────────────────────────────
    if data.get("skills"):
        section("Skills")
        for category, skill_list in data["skills"].items():
            if skill_list:
                story.append(Paragraph(
                    f"<b>{category}:</b> {', '.join(skill_list)}",
                    body_style
                ))

    doc.build(story)
    return buffer.getvalue()
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io as _io

def generate_resume_docx(data: dict) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    accent = RGBColor(0x2E, 0x7D, 0x52)
    dark = RGBColor(0x1C, 0x2C, 0x25)

    p = data.get("personal", {})

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(p.get("name", ""))
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = dark

    # Contact
    contact_parts = [x for x in [p.get("email",""), p.get("phone",""), p.get("linkedin",""), p.get("location","")] if x.strip()]
    contact_para = doc.add_paragraph("  |  ".join(contact_parts))
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.runs[0].font.size = Pt(9)

    def add_section(title):
        hp = doc.add_paragraph()
        run = hp.add_run(title.upper())
        run.bold = True
        run.font.color.rgb = accent
        run.font.size = Pt(10)
        doc.add_paragraph("─" * 60).runs[0].font.size = Pt(7)

    def add_bullet(text):
        bp = doc.add_paragraph(style="List Bullet")
        bp.add_run(text).font.size = Pt(9)

    if data.get("summary"):
        add_section("Professional Summary")
        doc.add_paragraph(data["summary"]).runs[0].font.size = Pt(9)

    if data.get("education"):
        add_section("Education")
        for edu in data["education"]:
            ep = doc.add_paragraph()
            ep.add_run(f"{edu.get('degree','')} — {edu.get('institution','')}").bold = True
            doc.add_paragraph(f"{edu.get('year','')}  {'CGPA: ' + edu.get('gpa','') if edu.get('gpa') else ''}").runs[0].font.size = Pt(9)

    if data.get("experience"):
        add_section("Experience")
        for exp in data["experience"]:
            xp = doc.add_paragraph()
            xp.add_run(f"{exp.get('role','')} — {exp.get('company','')}").bold = True
            doc.add_paragraph(exp.get("duration","")).runs[0].font.size = Pt(9)
            for b in exp.get("bullets", []):
                add_bullet(b)

    if data.get("projects"):
        add_section("Projects")
        for proj in data["projects"]:
            pp = doc.add_paragraph()
            pp.add_run(proj.get("name","")).bold = True
            if proj.get("tech"):
                pp.add_run(f" — {proj.get('tech','')}")
            for b in proj.get("bullets", []):
                add_bullet(b)

    if data.get("skills"):
        add_section("Skills")
        for category, skill_list in data["skills"].items():
            if skill_list:
                sp = doc.add_paragraph()
                sp.add_run(f"{category}: ").bold = True
                sp.add_run(", ".join(skill_list)).font.size = Pt(9)

    buffer = _io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()